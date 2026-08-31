"""
文档加载器

支持多种文件格式的加载和智能文本切分。
"""

import pandas as pd
from pypdf import PdfReader
from docx import Document
from pathlib import Path
import re


def smart_text_segmentation(text, max_chunk_size=2000, min_chunk_size=500):
    """
    智能文本切分：保持语义完整性，控制处理时间

    Args:
        text: 原始文本
        max_chunk_size: 最大块大小（字符数）
        min_chunk_size: 最小块大小（字符数）

    Returns:
        切分后的文本块列表
    """
    # 1. 预处理：清理特殊符号，保留语义关系
    cleaned_text = re.sub(r'[^一-龥a-zA-Z0-9\s,，.。:：;；!！?？]', '', text)
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()

    if not cleaned_text:
        return []

    # 2. 按段落分割
    paragraphs = [p.strip() for p in cleaned_text.split('\n') if p.strip()]

    # 3. 如果文本较短，直接返回
    if len(cleaned_text) <= max_chunk_size:
        return [cleaned_text]

    # 4. 智能合并段落，构建语义块
    chunks = []
    current_chunk = ""

    for paragraph in paragraphs:
        if len(paragraph) > max_chunk_size:
            sentences = re.split(r'[。！？!?]', paragraph)
            sentences = [s.strip() for s in sentences if s.strip()]

            for sentence in sentences:
                if len(current_chunk) + len(sentence) + 1 > max_chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk)
                        current_chunk = ""

                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence

        elif len(current_chunk) + len(paragraph) + 1 > max_chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = paragraph
            else:
                chunks.append(paragraph)
                current_chunk = ""
        else:
            if current_chunk:
                current_chunk += " " + paragraph
            else:
                current_chunk = paragraph

    if current_chunk:
        chunks.append(current_chunk)

    # 5. 合并过小的块
    merged_chunks = []
    temp_chunk = ""

    for chunk in chunks:
        if len(temp_chunk) + len(chunk) + 1 <= max_chunk_size:
            if temp_chunk:
                temp_chunk += " " + chunk
            else:
                temp_chunk = chunk
        else:
            if temp_chunk:
                merged_chunks.append(temp_chunk)
                temp_chunk = chunk
            else:
                merged_chunks.append(chunk)

    if temp_chunk:
        merged_chunks.append(temp_chunk)

    # 6. 确保每个块都有足够的语义内容
    final_chunks = []
    for chunk in merged_chunks:
        if len(chunk) >= min_chunk_size:
            final_chunks.append(chunk)
        else:
            if final_chunks:
                final_chunks[-1] += " " + chunk
            else:
                final_chunks.append(chunk)

    return final_chunks


def load_document(file_source, max_chunk_size=2000, min_chunk_size=500):
    """
    根据文件类型加载内容，返回智能切分的文本块列表

    Args:
        file_source: 文件路径字符串或上传的文件对象
        max_chunk_size: 最大块大小（字符数）
        min_chunk_size: 最小块大小（字符数）

    Returns:
        (文本块列表，错误信息)
    """
    if isinstance(file_source, str):
        file_path = Path(file_source)
        if not file_path.exists():
            return None, f"文件不存在：{file_source}"
        file_type = file_path.suffix.lower().lstrip('.')
    else:
        # Streamlit 上传的文件对象
        file_type = file_source.name.split('.')[-1].lower()

    try:
        text_content = _read_content(file_source, file_type)
    except Exception as e:
        return None, f"解析失败：{str(e)}"

    if not text_content:
        return None, "文档内容为空或无法解析"

    chunks = smart_text_segmentation(text_content, max_chunk_size, min_chunk_size)
    if not chunks:
        return None, "文档内容为空或无法解析"

    return chunks, None


def _read_content(source, file_type: str) -> str:
    """
    统一读取函数：支持文件路径（str/Path）和 Streamlit 上传文件对象。

    Args:
        source: 文件路径或上传文件对象
        file_type: 文件类型（不含点号，小写）

    Returns:
        提取的文本内容
    """
    if file_type in ('xlsx', 'xls'):
        return _read_excel(source)

    if file_type == 'pdf':
        return _read_pdf(source)

    if file_type in ('docx', 'doc'):
        return _read_docx(source)

    if file_type == 'txt':
        return _read_txt(source)

    return ""


def _read_excel(source):
    """读取 Excel 文件"""
    df = pd.read_excel(source)
    lines = []
    for _, row in df.iterrows():
        row_text = " ".join(str(cell) for cell in row if pd.notna(cell))
        if row_text.strip():
            lines.append(row_text)
    return "\n".join(lines)


def _read_pdf(source):
    """读取 PDF 文件"""
    if isinstance(source, (str, Path)):
        reader = PdfReader(str(source))
    else:
        reader = PdfReader(source)
    parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text and page_text.strip():
            parts.append(page_text)
    return "\n".join(parts)


def _read_docx(source):
    """读取 Word 文档"""
    if isinstance(source, (str, Path)):
        doc = Document(str(source))
    else:
        doc = Document(source)
    parts = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text)
    return "\n".join(parts)


def _read_txt(source):
    """读取文本文件（支持多种编码）"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']
    if isinstance(source, (str, Path)):
        for enc in encodings:
            try:
                with open(source, 'r', encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        return ""

    # Streamlit 上传文件对象
    raw_bytes = source.read()
    for enc in encodings:
        try:
            return raw_bytes.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return ""
